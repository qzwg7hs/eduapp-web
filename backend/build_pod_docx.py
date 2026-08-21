"""
Builds a sample "Problem of the Day" bulk-upload .docx matching the format
expected by app/pod_parser.py: each problem block has BOTH languages in one
entry (Kazakh question first, then Russian), an optional single shared image,
and a shared answer. Several problems in one file -> queued onto consecutive
days in file order.

Run from the backend/ directory with the venv active:
    venv\\Scripts\\python.exe build_pod_docx.py
"""
import struct
import zlib

from docx import Document
from docx.shared import Inches

OUT_PATH = r"C:\Users\Aruay\Desktop\eduapp-web\eduapp-web\pod_sample.docx"


def make_blank_png(size: int = 256) -> bytes:
    """A minimal solid-color PNG, built with only stdlib (no Pillow needed)."""
    def chunk(tag: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", size, size, 8, 2, 0, 0, 0)  # 8-bit, RGB
    raw_row = b"\x00" + b"\x00\x40\x80" * size  # filter=none, teal-ish pixels
    raw = raw_row * size
    idat = zlib.compress(raw, 9)
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


BLANK_IMAGE = make_blank_png()

# Six sample problems: a mix of with/without images, plain arithmetic and
# LaTeX-flavored text, matching the grade-6 curriculum already in the app.
PROBLEMS = [
    {
        "kz": "Пете дүкенге барды. Ол 3 дәптер сатып алды, әр дәптердің бағасы 150 теңге. Пете дүкенге қанша теңге төледі?",
        "ru": "Петя пошёл в магазин. Он купил 3 тетради, каждая тетрадь стоит 150 тенге. Сколько тенге заплатил Петя?",
        "answer": "450",
        "image": False,
    },
    {
        "kz": "Тіктөртбұрыштың ұзындығы 12 см, ені 5 см. Тіктөртбұрыштың периметрін тап.",
        "ru": "Длина прямоугольника 12 см, ширина 5 см. Найди периметр прямоугольника.",
        "answer": "34",
        "image": True,
    },
    {
        "kz": "$45\\%$-ы 60-қа тең сан қандай сан? Толық санмен жауап бер.",
        "ru": "Число, $45\\%$ которого равны 60 — какое это число? Ответь целым числом.",
        "answer": "133",
        "image": False,
    },
    {
        "kz": "Суретте көрсетілген үшбұрыштың периметрін есепте (қабырғалары: 7 см, 9 см, 11 см).",
        "ru": "Вычисли периметр треугольника, показанного на рисунке (стороны: 7 см, 9 см, 11 см).",
        "answer": "27",
        "image": True,
    },
    {
        "kz": "Алма мен алмұрттың қатынасы $3:4$. Егер 12 алма болса, алмұрт неше дана?",
        "ru": "Отношение яблок к грушам равно $3:4$. Если яблок 12 штук, сколько груш?",
        "answer": "16",
        "image": False,
    },
    {
        "kz": "Бір сыныпта 28 оқушы бар. Оның $\\frac{3}{7}$ бөлігі қыздар. Сыныпта неше қыз бар?",
        "ru": "В классе 28 учеников. $\\frac{3}{7}$ из них — девочки. Сколько девочек в классе?",
        "answer": "12",
        "image": False,
    },
]

doc = Document()

for i, p in enumerate(PROBLEMS, start=1):
    doc.add_paragraph(f"[ЗАДАЧА {i}]")
    doc.add_paragraph("[ҚАЗАҚША]")
    doc.add_paragraph(p["kz"])
    doc.add_paragraph("[РУССКИЙ]")
    doc.add_paragraph(p["ru"])
    if p["image"]:
        doc.add_paragraph("[ИЗОБРАЖЕНИЕ]")
        pic_para = doc.add_paragraph()
        run = pic_para.add_run()
        run.add_picture(io_bytes := __import__("io").BytesIO(BLANK_IMAGE), width=Inches(1.5))
    doc.add_paragraph("[ОТВЕТ]")
    doc.add_paragraph(p["answer"])
    doc.add_paragraph("")  # spacer between problems

doc.save(OUT_PATH)
print(f"Saved {len(PROBLEMS)} problems to {OUT_PATH}")
print(f"  with image: {sum(1 for p in PROBLEMS if p['image'])}, without: {sum(1 for p in PROBLEMS if not p['image'])}")
