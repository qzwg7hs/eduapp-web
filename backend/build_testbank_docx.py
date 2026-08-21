"""
One-off script: merges the 4 drafted JSON batches (numbers 1-100, bilingual)
into two .docx files (kz and ru) matching the bracket-tag format expected by
app/test_parser.py, ready to be uploaded through the real admin endpoint.

Run from the backend/ directory with the venv active:
    venv\\Scripts\\python.exe build_testbank_docx.py
"""
import io
import json
import struct
import zlib

from docx import Document
from docx.shared import Inches

BATCH_DIR = r"C:\Users\Aruay\AppData\Local\Temp\claude\c--Users-Aruay-Desktop-eduapp-web\e30c8997-15e3-44e3-9bf1-080a0ec2a20b\scratchpad\testbank"
BATCH_FILES = ["batch_1_25.json", "batch_26_50.json", "batch_51_75.json", "batch_76_100.json"]

OUT_KZ = r"C:\Users\Aruay\Desktop\eduapp-web\eduapp-web\testbank_kz.docx"
OUT_RU = r"C:\Users\Aruay\Desktop\eduapp-web\eduapp-web\testbank_ru.docx"

LETTERS = ["A", "B", "C", "D", "E", "F"]

# Question numbers to attach a sample image to, purely to demonstrate the
# [ИЗОБРАЖЕНИЕ] tag mechanics — one MCQ (#1) and one open question (#2).
DEMO_IMAGE_NUMBERS = {1, 2}


def make_blank_png(size: int = 256) -> bytes:
    """A minimal solid-black PNG, built with only stdlib (no Pillow needed)."""
    def chunk(tag: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", size, size, 8, 2, 0, 0, 0)  # 8-bit, RGB
    raw_row = b"\x00" + b"\x00\x00\x00" * size  # filter=none, all-black pixels
    raw = raw_row * size
    idat = zlib.compress(raw, 9)
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


BLANK_IMAGE = make_blank_png()

questions = []
for fname in BATCH_FILES:
    with open(f"{BATCH_DIR}\\{fname}", encoding="utf-8") as f:
        data = json.load(f)
    questions.extend(data["questions"])

questions.sort(key=lambda q: q["number"])

# Sanity check: numbers 1..100 each exactly once
numbers = [q["number"] for q in questions]
assert numbers == list(range(1, 101)), f"Number sequence broken: {numbers[:5]}...{numbers[-5:]} (count={len(numbers)})"

mcq_count = sum(1 for q in questions if q["kz"]["problem_type"] == "mcq")
open_count = len(questions) - mcq_count
print(f"Loaded {len(questions)} questions ({mcq_count} mcq, {open_count} open)")


def build_docx(lang: str, out_path: str):
    doc = Document()
    for q in questions:
        item = q[lang]
        doc.add_paragraph(f"[ЗАДАЧА {q['number']}]")
        doc.add_paragraph("[ОПИСАНИЕ]")
        doc.add_paragraph(item["question"])
        if q["number"] in DEMO_IMAGE_NUMBERS:
            doc.add_paragraph("[ИЗОБРАЖЕНИЕ]")
            doc.add_picture(io.BytesIO(BLANK_IMAGE), width=Inches(1.5))
        doc.add_paragraph("[ТИП ОТВЕТА]")
        if item["problem_type"] == "mcq":
            doc.add_paragraph("Тест с вариантами")
            doc.add_paragraph("[ВАРИАНТЫ ОТВЕТА]")
            for i, opt in enumerate(item["options"]):
                doc.add_paragraph(f"{LETTERS[i]}) {opt}")
            doc.add_paragraph("[ОТВЕТ]")
            doc.add_paragraph(LETTERS[item["correct_index"]])
        else:
            doc.add_paragraph("Одно число")
            doc.add_paragraph("[ОТВЕТ]")
            doc.add_paragraph(str(item["answer"]))
    doc.save(out_path)
    print(f"Wrote {out_path}")


build_docx("kz", OUT_KZ)
build_docx("ru", OUT_RU)
print("Done.")
